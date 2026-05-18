"""
exports.py — Generate PDF and Word (.docx) exports for research papers.

Single paper:  generate_pdf(paper_dict)  → bytes
               generate_docx(paper_dict) → bytes

All papers:    generate_all_pdf(papers)  → bytes
               generate_all_docx(papers) → bytes
"""

import io
from datetime import datetime
from typing import Optional

# ── PDF via fpdf2 ─────────────────────────────────────────────────────────────
from fpdf import FPDF, XPos, YPos

BLUE  = (30, 77, 145)
LBLUE = (46, 117, 182)
LGRAY = (245, 247, 251)
DGRAY = (100, 100, 110)
BLACK = (26, 26, 46)
WHITE = (255, 255, 255)
HIGH_C  = (192, 57, 43)
MED_C   = (180, 140, 10)
LOW_C   = (30, 132, 73)


def _vleo_color(rel: str):
    if "High"   in rel: return HIGH_C
    if "Medium" in rel: return MED_C
    if "Low"    in rel: return LOW_C
    return DGRAY


class PaperPDF(FPDF):
    def header(self):
        self.set_fill_color(*BLUE)
        self.rect(0, 0, 210, 12, "F")
        self.set_font("Helvetica", "B", 9)
        self.set_text_color(*WHITE)
        self.set_xy(10, 2)
        self.cell(0, 8, "VLEO Research Papers Database", align="L")
        self.set_xy(0, 2)
        self.cell(200, 8, datetime.now().strftime("%d %b %Y"), align="R")
        self.set_text_color(*BLACK)
        self.ln(10)

    def footer(self):
        self.set_y(-12)
        self.set_font("Helvetica", "", 8)
        self.set_text_color(*DGRAY)
        self.cell(0, 8, f"Page {self.page_no()}", align="C")
        self.set_text_color(*BLACK)


def _section_title(pdf: FPDF, title: str):
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*LBLUE)
    pdf.set_fill_color(*LGRAY)
    pdf.cell(0, 7, f"  {title}", new_x=XPos.LMARGIN, new_y=YPos.NEXT, fill=True)
    pdf.set_text_color(*BLACK)
    pdf.ln(1)


def _body_text(pdf: FPDF, text: str, indent: float = 5):
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(*BLACK)
    pdf.set_x(pdf.l_margin + indent)
    pdf.multi_cell(0, 5.5, text or "—", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(3)


def _vleo_box(pdf: FPDF, notes: dict, overall: str):
    if not notes:
        return
    criteria = [
        ("altitude_coverage",         "Altitude coverage (<450 km)"),
        ("atmospheric_drag",          "Atmospheric drag & orbital decay"),
        ("satellite_design_for_vleo", "Satellite design for VLEO"),
        ("collision_debris_risk",     "Collision & debris risk at VLEO"),
    ]
    col_w = 92
    x0 = pdf.l_margin

    for i, (key, label) in enumerate(criteria):
        score = notes.get(key, "❌")
        col = i % 2
        if col == 0:
            pdf.set_x(x0)
        else:
            pdf.set_x(x0 + col_w + 4)

        pdf.set_fill_color(*LGRAY)
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(8, 7, score, fill=True)
        pdf.set_font("Helvetica", "", 9)
        pdf.set_fill_color(235, 240, 250)
        pdf.cell(col_w - 8, 7, f"  {label}", fill=True,
                 new_x=XPos.RIGHT if col == 0 else XPos.LMARGIN,
                 new_y=YPos.LAST if col == 0 else YPos.NEXT)
        if col == 1:
            pdf.ln(1)

    pdf.ln(3)
    # Verdict
    verdict = notes.get("verdict", "")
    color = _vleo_color(overall)
    pdf.set_fill_color(*color)
    pdf.set_text_color(*WHITE)
    pdf.set_font("Helvetica", "B", 10)
    label_txt = overall.replace("🔴","").replace("🟡","").replace("🟢","").replace("⬜","").strip()
    pdf.cell(30, 7, f"  {label_txt}", fill=True)
    pdf.set_fill_color(245, 245, 245)
    pdf.set_text_color(*BLACK)
    pdf.set_font("Helvetica", "I", 9)
    pdf.multi_cell(0, 7, f"  {verdict}", fill=True,
                   new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(3)


def _render_paper_pdf(pdf: FPDF, p: dict, include_header: bool = True):
    """Render one paper into the PDF object."""
    # Title block
    pdf.set_fill_color(*BLUE)
    pdf.set_text_color(*WHITE)
    pdf.set_font("Helvetica", "B", 14)
    pdf.multi_cell(0, 8, p.get("title", "Untitled"),
                   fill=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_text_color(*WHITE)
    pdf.set_font("Helvetica", "", 9)
    meta = " · ".join(filter(None, [
        p.get("authors"), str(p.get("year") or ""), p.get("journal_conference")
    ]))
    pdf.set_fill_color(*LBLUE)
    pdf.cell(0, 6, f"  {meta}", fill=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_text_color(*BLACK)
    pdf.ln(4)

    # Chips row
    chips = [p.get("vleo_relevance"), p.get("altitude_range"),
             p.get("methodology_type"), p.get("mission_type")]
    chips = [c for c in chips if c]
    if chips:
        for chip in chips:
            color = _vleo_color(chip) if "High" in chip or "Medium" in chip or "Low" in chip else LBLUE
            pdf.set_fill_color(*color)
            pdf.set_text_color(*WHITE)
            pdf.set_font("Helvetica", "B", 8)
            clean = chip.replace("🔴","").replace("🟡","").replace("🟢","").replace("⬜","").strip()
            pdf.cell(len(clean)*2.3 + 6, 6, f" {clean} ", fill=True)
            pdf.cell(2, 6, "")
        pdf.set_text_color(*BLACK)
        pdf.ln(8)

    # Tags
    tags = p.get("tags", "")
    if tags:
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(*DGRAY)
        pdf.cell(0, 5, "Tags: " + tags, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(*BLACK)
        pdf.ln(3)

    # Content sections
    fields = [
        ("Aim / Research Question", "aim"),
        ("Methodology", "methodology"),
        ("Key Results", "key_results"),
        ("Figures & Graphs", "figures_graphs"),
        ("Summary", "summary"),
        ("References", "references"),
    ]
    for label, key in fields:
        val = p.get(key)
        if val:
            _section_title(pdf, label)
            _body_text(pdf, val)

    # VLEO assessment
    notes = p.get("vleo_relevance_notes")
    overall = p.get("vleo_relevance", "")
    if notes and overall:
        _section_title(pdf, "VLEO Relevance Assessment")
        _vleo_box(pdf, notes, overall)

    # Links
    doi = p.get("doi_url")
    notion = p.get("notion_url")
    if doi or notion:
        _section_title(pdf, "Links")
        if doi:
            _body_text(pdf, f"Paper: {doi}")
        if notion:
            _body_text(pdf, f"Notion: {notion}")


def generate_pdf(paper: dict) -> bytes:
    """Generate a PDF for a single paper. Returns bytes."""
    pdf = PaperPDF()
    pdf.set_margins(15, 18, 15)
    pdf.set_auto_page_break(True, margin=18)
    pdf.add_page()
    _render_paper_pdf(pdf, paper)
    return bytes(pdf.output())


def generate_all_pdf(papers: list[dict]) -> bytes:
    """Generate a combined PDF report for all papers."""
    pdf = PaperPDF()
    pdf.set_margins(15, 18, 15)
    pdf.set_auto_page_break(True, margin=18)

    # Cover page
    pdf.add_page()
    pdf.set_fill_color(*BLUE)
    pdf.rect(0, 0, 210, 297, "F")
    pdf.set_text_color(*WHITE)
    pdf.set_y(80)
    pdf.set_font("Helvetica", "B", 28)
    pdf.cell(0, 14, "Research Papers", align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 10, "VLEO Database Export", align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(10)
    pdf.set_font("Helvetica", "", 13)
    pdf.cell(0, 8, f"{len(papers)} papers · Exported {datetime.now().strftime('%d %b %Y')}",
             align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # Stats
    high = sum(1 for p in papers if "High" in (p.get("vleo_relevance") or ""))
    med  = sum(1 for p in papers if "Medium" in (p.get("vleo_relevance") or ""))
    low  = sum(1 for p in papers if "Low" in (p.get("vleo_relevance") or "") and "Medium" not in (p.get("vleo_relevance") or ""))
    pdf.ln(20)
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, f"🔴 High: {high}   🟡 Medium: {med}   🟢 Low: {low}", align="C")
    pdf.set_text_color(*BLACK)

    # Each paper on its own page
    for p in papers:
        pdf.add_page()
        _render_paper_pdf(pdf, p)

    return bytes(pdf.output())


# ── DOCX via python-docx ──────────────────────────────────────────────────────
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_run_color(run, rgb: tuple):
    run.font.color.rgb = RGBColor(*rgb)


def _add_heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor(*BLUE) if level == 1 else RGBColor(*LBLUE)
    return p


def _add_field(doc, label: str, value: str):
    if not value:
        return
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.color.rgb = RGBColor(*LBLUE)
    run_val = p.add_run(value)
    run_val.font.color.rgb = RGBColor(*BLACK)
    p.paragraph_format.space_after = Pt(6)


def _render_paper_docx(doc, p: dict, page_break: bool = False):
    if page_break:
        doc.add_page_break()

    # Title
    title_p = doc.add_heading(p.get("title", "Untitled"), level=1)
    title_p.runs[0].font.color.rgb = RGBColor(*BLUE)
    title_p.runs[0].font.size = Pt(16)

    # Meta line
    meta = " · ".join(filter(None, [
        p.get("authors"), str(p.get("year") or ""), p.get("journal_conference")
    ]))
    meta_p = doc.add_paragraph(meta)
    meta_p.runs[0].italic = True
    meta_p.runs[0].font.color.rgb = RGBColor(*DGRAY)
    meta_p.paragraph_format.space_after = Pt(8)

    # Classification row
    chips = [p.get("vleo_relevance"), p.get("altitude_range"),
             p.get("methodology_type"), p.get("mission_type")]
    chips = [c for c in chips if c]
    if chips:
        chip_p = doc.add_paragraph()
        for chip in chips:
            clean = chip.replace("🔴","").replace("🟡","").replace("🟢","").replace("⬜","").strip()
            r = chip_p.add_run(f" {clean} ")
            r.bold = True
            r.font.size = Pt(9)
            chip_p.add_run("  ")
        chip_p.paragraph_format.space_after = Pt(8)

    # Tags
    tags = p.get("tags")
    if tags:
        _add_field(doc, "Tags", tags)

    # Horizontal rule
    hr = doc.add_paragraph()
    hr.paragraph_format.space_after = Pt(4)

    # Content sections
    sections = [
        ("Aim / Research Question", "aim"),
        ("Methodology", "methodology"),
        ("Key Results", "key_results"),
        ("Figures & Graphs", "figures_graphs"),
        ("Summary", "summary"),
        ("References", "references"),
    ]
    for label, key in sections:
        val = p.get(key)
        if val:
            h = doc.add_heading(label, level=2)
            h.runs[0].font.color.rgb = RGBColor(*LBLUE)
            body = doc.add_paragraph(val)
            body.paragraph_format.space_after = Pt(8)

    # VLEO section
    notes = p.get("vleo_relevance_notes")
    overall = p.get("vleo_relevance", "")
    if notes and overall:
        h = doc.add_heading("VLEO Relevance Assessment", level=2)
        h.runs[0].font.color.rgb = RGBColor(*LBLUE)
        criteria = [
            ("altitude_coverage",         "Altitude coverage (<450 km)"),
            ("atmospheric_drag",          "Atmospheric drag & orbital decay"),
            ("satellite_design_for_vleo", "Satellite design for VLEO"),
            ("collision_debris_risk",     "Collision & debris risk at VLEO"),
        ]
        for key, label in criteria:
            score = notes.get(key, "❌")
            cp = doc.add_paragraph()
            r1 = cp.add_run(f"{score}  ")
            r1.font.size = Pt(12)
            r2 = cp.add_run(label)
            r2.font.size = Pt(10)
            cp.paragraph_format.space_after = Pt(3)

        verdict = notes.get("verdict", "")
        if verdict:
            vp = doc.add_paragraph()
            rl = vp.add_run("Verdict: ")
            rl.bold = True
            color = _vleo_color(overall)
            rl.font.color.rgb = RGBColor(*color)
            rv = vp.add_run(verdict)
            rv.italic = True
            vp.paragraph_format.space_after = Pt(8)

    # Links
    doi = p.get("doi_url")
    notion = p.get("notion_url")
    if doi or notion:
        doc.add_heading("Links", level=2)
        if doi:
            _add_field(doc, "Paper", doi)
        if notion:
            _add_field(doc, "Notion", notion)


def generate_docx(paper: dict) -> bytes:
    """Generate a Word doc for a single paper. Returns bytes."""
    doc = DocxDocument()
    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    _render_paper_docx(doc, paper, page_break=False)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_all_docx(papers: list[dict]) -> bytes:
    """Generate a combined Word doc for all papers."""
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Cover
    cover = doc.add_heading("Research Papers — VLEO Database", level=1)
    cover.runs[0].font.color.rgb = RGBColor(*BLUE)
    cover.runs[0].font.size = Pt(24)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        f"{len(papers)} papers · Exported {datetime.now().strftime('%d %b %Y')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(*DGRAY)
    sub.paragraph_format.space_after = Pt(24)

    high = sum(1 for p in papers if "High" in (p.get("vleo_relevance") or ""))
    med  = sum(1 for p in papers if "Medium" in (p.get("vleo_relevance") or ""))
    low  = sum(1 for p in papers if "Low" in (p.get("vleo_relevance") or "") and "Medium" not in (p.get("vleo_relevance") or ""))
    stats = doc.add_paragraph(f"🔴 High Relevance: {high}    🟡 Medium: {med}    🟢 Low: {low}")
    stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stats.paragraph_format.space_after = Pt(12)

    for i, p in enumerate(papers):
        _render_paper_docx(doc, p, page_break=True)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
